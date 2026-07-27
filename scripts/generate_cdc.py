#!/usr/bin/env python3
"""Génère le cahier des charges Word du site vitrine BTP offshore."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "Cahier_des_charges_Site_Vitrine_BTP_Offshore.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(1.25 * (level + 1))
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = hdr[i]._element.get_or_add_tcPr()
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1A3A5C")
        shd.set(qn("w:val"), "clear")
        shading.append(shd)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # --- Page de garde ---
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CAHIER DES CHARGES")
    set_run_font(run, size=28, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Site vitrine — Entreprise BTP Offshore\nPlans & Conception architecturale")
    set_run_font(run, size=16, bold=False, color=RGBColor(0x3D, 0x5A, 0x74))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "\nVersion 1.0 — Juillet 2026\n"
        "Document de spécification fonctionnelle et technique\n"
        "Public cible principal : France (clients, partenaires, donneurs d’ordre)"
    )
    set_run_font(run, size=11, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()

    # --- Sommaire ---
    add_heading_custom(doc, "Sommaire", 1)
    toc_items = [
        "1. Contexte et présentation du projet",
        "2. Objectifs du site",
        "3. Public cible et positionnement",
        "4. Périmètre fonctionnel",
        "5. Expérience 3D scroll-driven (cœur du produit)",
        "6. Architecture de contenu (sections)",
        "7. Identité visuelle et UX",
        "8. Architecture technique recommandée",
        "9. Stack détaillée et justifications",
        "10. Exigences non fonctionnelles",
        "11. Analyse objective et risques",
        "12. Améliorations et suggestions",
        "13. Livrables et jalons",
        "14. Critères d’acceptation",
        "15. Annexes",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=4)

    doc.add_page_break()

    # --- 1 ---
    add_heading_custom(doc, "1. Contexte et présentation du projet", 1)
    add_para(
        doc,
        "Le présent cahier des charges définit la conception et le développement d’un site "
        "vitrine pour une activité indépendante / petite entreprise spécialisée en BTP offshore "
        "(conception de plans, études et prestations liées au bâtiment / offshore).",
    )
    add_para(doc, "Profil professionnel du porteur de projet :", bold=True)
    add_bullet(doc, "15 ans d’expérience en entreprise (BTP / plans / conception).")
    add_bullet(doc, "20 ans d’activité en freelance.")
    add_bullet(doc, "Collaboration régulière avec des entreprises étrangères.")
    add_bullet(
        doc,
        "Offre : tous types de plans (architecture, techniques, exécution, selon le périmètre métier à préciser).",
    )
    add_para(
        doc,
        "Le site a pour vocation d’être une vitrine crédible, moderne et mémorable, orientée "
        "acquisition de contacts (France en priorité), en valorisant le portfolio, l’expérience "
        "et les canaux de contact (LinkedIn, e-mail, Facebook, et éventuellement WhatsApp / téléphone).",
    )

    # --- 2 ---
    add_heading_custom(doc, "2. Objectifs du site", 1)
    add_heading_custom(doc, "2.1 Objectifs business", 2)
    add_bullet(doc, "Générer des prises de contact qualifiées (devis, missions, partenariats).")
    add_bullet(doc, "Renforcer la crédibilité auprès de clients et entreprises en France.")
    add_bullet(doc, "Différencier l’offre via une expérience 3D immersive (effet « wow », mémorabilité).")
    add_bullet(doc, "Présenter clairement le parcours, les compétences et des projets représentatifs.")

    add_heading_custom(doc, "2.2 Objectifs produit", 2)
    add_bullet(doc, "Une page (ou un parcours) scroll unique immersif, avec narration visuelle 3D.")
    add_bullet(doc, "Contenu commercial lisible pendant le scroll (expérience, preuves, projets).")
    add_bullet(doc, "Section contact claire, accessible, avec disparition dynamique de la maison 3D.")
    add_bullet(doc, "Performance et accessibilité suffisantes pour ne pas perdre les visiteurs moins techniques.")

    # --- 3 ---
    add_heading_custom(doc, "3. Public cible et positionnement", 1)
    add_table(
        doc,
        ["Segment", "Besoin", "Message clé"],
        [
            (
                "Maîtres d’ouvrage / particuliers (FR)",
                "Confiance, clarté, contact simple",
                "Expertise long terme, plans de qualité, réactivité",
            ),
            (
                "Entreprises BTP / bureaux d’études (FR)",
                "Compétence, collaboration, délais",
                "15+20 ans d’expérience, collab. internationales",
            ),
            (
                "Partenaires étrangers",
                "Références, capacité offshore",
                "Habitué aux collaborations cross-border",
            ),
        ],
    )
    add_para(
        doc,
        "Positionnement recommandé : « Expert plans & conception BTP, vitrine immersive — "
        "sérieux métier + modernité technique ». Éviter un site trop « jeu vidéo » au détriment "
        "de la clarté commerciale.",
        italic=True,
    )

    # --- 4 ---
    add_heading_custom(doc, "4. Périmètre fonctionnel", 1)
    add_heading_custom(doc, "4.1 Inclus (MVP)", 2)
    add_bullet(doc, "Landing / parcours scroll unique immersif.")
    add_bullet(doc, "Scène 3D maison (façade → zoom → porte → visite intérieure → 1er étage).")
    add_bullet(doc, "Narration textuelle synchronisée au scroll (présentation, expérience, projets).")
    add_bullet(doc, "Galerie / vitrine de projets (sélection 4–8 projets max en MVP).")
    add_bullet(doc, "Bloc expérience / parcours professionnel.")
    add_bullet(doc, "Section Contact (LinkedIn, e-mail, Facebook ; option téléphone / WhatsApp).")
    add_bullet(doc, "Disparition / transition de la maison 3D en bas de page (section Contact).")
    add_bullet(doc, "Responsive desktop + mobile (avec mode dégradé sur mobile si nécessaire).")
    add_bullet(doc, "Mentions légales / politique de confidentialité (RGPD — France).")

    add_heading_custom(doc, "4.2 Hors périmètre (V1)", 2)
    add_bullet(doc, "Espace client / login.")
    add_bullet(doc, "Devis en ligne automatisé / paiement.")
    add_bullet(doc, "CMS complet multi-auteurs (sauf besoin admin simple projets).")
    add_bullet(doc, "Blog éditorial (peut être une phase 2).")
    add_bullet(doc, "Multilingue complet (FR d’abord ; EN en phase 2 recommandé).")

    add_heading_custom(doc, "4.3 Canaux de contact", 2)
    add_table(
        doc,
        ["Canal", "Priorité", "Implémentation"],
        [
            ("E-mail", "P0", "Lien mailto + formulaire optionnel (recommandé)"),
            ("LinkedIn", "P0", "Lien profil / bouton CTA"),
            ("Facebook", "P1", "Lien page / profil"),
            ("Téléphone", "P1", "Lien tel: (si souhaité)"),
            ("WhatsApp", "P2", "Lien wa.me (fort pour international)"),
        ],
    )

    # --- 5 ---
    add_heading_custom(doc, "5. Expérience 3D scroll-driven (cœur du produit)", 1)
    add_para(
        doc,
        "Le site s’articule autour d’une maison 3D pilotée par le scroll. L’utilisateur progresse "
        "dans le récit en descendant la page ; la caméra et les animations de scène évoluent en conséquence.",
    )

    add_heading_custom(doc, "5.1 Scénario narratif (storyboard)", 2)
    add_numbered(
        doc,
        "Entrée : vue façade principale de la maison (plan large, branding + accroche courte).",
    )
    add_numbered(
        doc,
        "Scroll 1 — Approche : zoom progressif vers la façade (contenu : qui je suis / positionnement).",
    )
    add_numbered(
        doc,
        "Scroll 2 — Seuil : approche de la porte ; textes expérience (15 ans entreprise + 20 ans freelance).",
    )
    add_numbered(
        doc,
        "Ouverture porte : animation déclenchée (seuil de scroll). Une fois ouverte, une séquence "
        "autonome peut continuer même sans scroll supplémentaire (visite guidée).",
    )
    add_numbered(
        doc,
        "Visite intérieure : circulation RDC → montée / transition vers le 1er étage (projets / preuves).",
    )
    add_numbered(
        doc,
        "Fin de parcours / Contact : la maison disparaît dynamiquement (fade / dissolve / sortie de caméra) "
        "pour laisser place à un bloc Contact lisible et calme.",
    )

    add_heading_custom(doc, "5.2 Règles d’interaction", 2)
    add_bullet(
        doc,
        "Scroll progressif : plus l’utilisateur scrolle, plus la caméra avance / zoome (mapping scroll → timeline).",
    )
    add_bullet(
        doc,
        "Après ouverture de la porte : mode « auto-tour » possible (animation continue) tout en permettant "
        "au scroll de reprendre le contrôle sur les étapes suivantes.",
    )
    add_bullet(
        doc,
        "Pendant tout le parcours : overlays textuels (titres, phrases courtes, CTAs discrets) sans masquer la 3D.",
    )
    add_bullet(
        doc,
        "Section Contact : transition qui retire la scène 3D (ou la passe en arrière-plan très atténué) "
        "pour prioriser la conversion.",
    )
    add_bullet(
        doc,
        "Contrôles utilisateur : bouton « Passer l’intro / Aller au contact » obligatoire (accessibilité & impatience).",
    )

    add_heading_custom(doc, "5.3 Actifs 3D", 2)
    add_bullet(doc, "Modèle maison optimisé (GLB/GLTF), low-poly ou LOD, textures compressées (KTX2/Basis si possible).")
    add_bullet(doc, "Éclairage réaliste mais léger (pas de path-tracing temps réel).")
    add_bullet(doc, "Budget perf cible desktop : ≥ 45–60 FPS sur machine moyenne.")
    add_bullet(doc, "Fallback : image / vidéo courte ou version 2D si WebGL indisponible.")

    # --- 6 ---
    add_heading_custom(doc, "6. Architecture de contenu (sections)", 1)
    add_table(
        doc,
        ["Section", "Objectif", "Contenu type"],
        [
            (
                "Hero / Façade",
                "Ancrage marque + promesse",
                "Nom / activité, 1 phrase, CTA Contact / Projets",
            ),
            (
                "Présentation",
                "Se vendre clairement",
                "Expertise plans, offshore, collab. internationales",
            ),
            (
                "Expérience",
                "Preuve de sérieux",
                "15 ans entreprise, 20 ans freelance, faits chiffrés",
            ),
            (
                "Projets",
                "Preuve visuelle",
                "4–8 projets : image, type de plan, année, rôle",
            ),
            (
                "Méthode / Offre",
                "Clarifier l’offre",
                "Types de plans, process, zones d’intervention",
            ),
            (
                "Contact",
                "Conversion",
                "Mail, LinkedIn, Facebook, formulaire optionnel",
            ),
        ],
    )
    add_para(
        doc,
        "Recommandation éditoriale : phrases courtes, français soigné, chiffres concrets. "
        "Éviter le jargon trop technique sans explication pour les particuliers.",
        italic=True,
    )

    # --- 7 ---
    add_heading_custom(doc, "7. Identité visuelle et UX", 1)
    add_bullet(doc, "Direction artistique : architecture / chantier / précision — tons sobres (bleu acier, béton, bois, blanc cassé).")
    add_bullet(doc, "Typographie expressive (pas Inter/Roboto/Arial par défaut) ; hiérarchie forte.")
    add_bullet(doc, "Une composition par viewport ; éviter l’effet « dashboard ».")
    add_bullet(doc, "Pas de cartes décoratives inutiles ; la 3D est l’ancre visuelle principale.")
    add_bullet(doc, "Motion : 2–3 intentions fortes (zoom caméra, porte, disparition finale) — pas d’animation bruit.")
    add_bullet(doc, "Contraste texte / fond suffisant (WCAG AA cible).")
    add_bullet(doc, "Mobile : expérience simplifiée (scroll + textes + galerie) si la 3D lourde dégrade trop.")

    # --- 8 ---
    add_heading_custom(doc, "8. Architecture technique recommandée", 1)
    add_para(
        doc,
        "Choix retenu pour ce dépôt : une architecture hybride pragmatique, centrée React pour "
        "l’expérience 3D/scroll, avec un backend Python léger uniquement si nécessaire (formulaire, "
        "admin projets, envoi d’e-mails).",
        bold=True,
    )

    add_heading_custom(doc, "8.1 Schéma cible", 2)
    add_para(
        doc,
        "Frontend : Next.js (React) + React Three Fiber (Three.js) + GSAP ScrollTrigger (ou équivalent) "
        "pour le mapping scroll → timeline caméra.\n"
        "Backend optionnel : FastAPI (Python) pour formulaire de contact, stockage métadonnées projets, "
        "proxy d’envoi e-mail.\n"
        "Hébergement : Vercel/Netlify (front) + Railway/Render/Fly.io (API Python) ou tout en monorepo "
        "avec API routes Next.js si le besoin backend reste minimal.",
    )

    add_heading_custom(doc, "8.2 Pourquoi ce choix (et pas « tout Python » ou « tout React » aveugle)", 2)
    add_bullet(
        doc,
        "La 3D Web et le scroll-driven storytelling sont nativement meilleurs en JavaScript/WebGL "
        "(Three.js / R3F). Django/Flask seuls ne sont pas adaptés pour piloter cette UX.",
    )
    add_bullet(
        doc,
        "Un site vitrine peut démarrer 100 % Next.js (SSG/SSR + API routes) sans Python. "
        "Python devient pertinent dès qu’on veut un admin métier, un pipeline d’assets, ou une API claire.",
    )
    add_bullet(
        doc,
        "Recommandation pragmatique MVP : Next.js + R3F + contenu en Markdown/JSON. "
        "Ajouter FastAPI en phase 2 si formulaire + back-office projets deviennent nécessaires.",
    )

    # --- 9 ---
    add_heading_custom(doc, "9. Stack détaillée et justifications", 1)
    add_table(
        doc,
        ["Couche", "Technologie", "Rôle"],
        [
            ("UI / App", "Next.js 15 (App Router) + TypeScript", "Pages, perf, SEO, structure"),
            ("3D", "Three.js + React Three Fiber + Drei", "Scène maison, caméra, assets"),
            ("Scroll / motion", "GSAP + ScrollTrigger (ou Lenis)", "Timeline liée au scroll"),
            ("Styles", "Tailwind CSS + CSS variables", "Design system léger"),
            ("Contenu projets", "MDX / JSON (MVP)", "Portfolio facilement éditable"),
            ("Formulaire (opt.)", "API Route Next ou FastAPI", "Contact + anti-spam"),
            ("E-mail (opt.)", "Resend / SMTP", "Notification des messages"),
            ("Analytics", "Plausible ou GA4 (consentement)", "Mesure des conversions"),
            ("CI / qualité", "ESLint, Prettier, Playwright smoke", "Stabilité des releases"),
        ],
    )

    add_heading_custom(doc, "9.1 Structure de dépôt proposée", 2)
    add_para(
        doc,
        "btp-site/\n"
        "  apps/web/          → Next.js (expérience 3D + pages)\n"
        "  apps/api/          → FastAPI (optionnel, phase 2)\n"
        "  packages/assets/   → modèles GLB, textures, images projets\n"
        "  docs/              → CDC, storyboard, assets briefs\n"
        "  README.md",
        size=10,
    )

    # --- 10 ---
    add_heading_custom(doc, "10. Exigences non fonctionnelles", 1)
    add_table(
        doc,
        ["Thème", "Exigence"],
        [
            ("Performance", "LCP < 2.5s hors scène 3D ; chargement 3D progressif (lazy + suspense)"),
            ("SEO", "Métadonnées FR, Open Graph, titres clairs ; contenu textuel crawlable"),
            ("Accessibilité", "Skip link, CTA Contact toujours accessible, prefers-reduced-motion"),
            ("Compatibilité", "Chrome/Edge/Firefox/Safari récents ; fallback WebGL"),
            ("Sécurité", "Pas de secrets côté client ; protection formulaire (honeypot/rate limit)"),
            ("RGPD", "Bandeau cookies si tracking ; pages légales FR"),
            ("Maintenabilité", "TypeScript strict, composants isolés, assets versionnés"),
        ],
    )

    # --- 11 ---
    add_heading_custom(doc, "11. Analyse objective et risques", 1)
    add_para(
        doc,
        "Le concept 3D scroll-driven est un excellent différenciateur, mais c’est aussi le principal "
        "risque projet : coût d’asset, perf mobile, et risque de « beau site peu clair » si le discours "
        "commercial est trop secondaire.",
        bold=True,
    )
    add_table(
        doc,
        ["Risque", "Impact", "Mitigation"],
        [
            (
                "Modèle 3D trop lourd",
                "Élevé",
                "Budget poly/textures ; LOD ; compression ; lazy load",
            ),
            (
                "Mobile lent / chauffe",
                "Élevé",
                "Mode dégradé mobile (vidéo/images + scroll 2D)",
            ),
            (
                "UX confuse (scroll vs auto-tour)",
                "Moyen",
                "Storyboard précis + skip + tests utilisateurs",
            ),
            (
                "Contenu projets insuffisant",
                "Élevé",
                "Sélectionner 4–8 projets forts avec droits d’image",
            ),
            (
                "SEO faible (site trop « canvas »)",
                "Moyen",
                "Texte HTML réel, titres, balises, sitemap",
            ),
            (
                "Dépendance à un seul effet wow",
                "Moyen",
                "CTA Contact visibles tôt ; preuves métier solides",
            ),
            (
                "Scope creep (trop d’animations)",
                "Moyen",
                "MVP strict : 1 maison, 1 timeline, 1 disparition",
            ),
        ],
    )

    # --- 12 ---
    add_heading_custom(doc, "12. Améliorations et suggestions", 1)
    add_para(
        doc,
        "Suggestions objectives pour maximiser le ROI du site (au-delà de la demande initiale) :",
    )

    add_heading_custom(doc, "12.1 Produit & conversion", 2)
    add_bullet(
        doc,
        "CTA Contact récurrent discret (header sticky ou bouton flottant) — ne pas attendre le bas de page.",
    )
    add_bullet(
        doc,
        "Formulaire de contact court (nom, e-mail, besoin, budget indicatif) en plus des liens sociaux.",
    )
    add_bullet(
        doc,
        "Preuves sociales : logos clients / pays collaborés / types de missions (même anonymisés).",
    )
    add_bullet(
        doc,
        "Page Projets séparée (liste + détail) pour le SEO, en plus de la vitrine dans le scroll.",
    )
    add_bullet(
        doc,
        "Version EN (phase 2) : utile vu les collaborations étrangères.",
    )

    add_heading_custom(doc, "12.2 Expérience 3D", 2)
    add_bullet(
        doc,
        "Prévoir un mode « reduced motion » : afficher une séquence d’images / vidéo pré-rendue.",
    )
    add_bullet(
        doc,
        "Sur mobile, remplacer la visite WebGL lourde par une vidéo scrollée (scrollytelling) ou un scrub vidéo.",
    )
    add_bullet(
        doc,
        "Ajouter un hotspot « Voir un plan 2D » pendant la visite (lien vers un plan PDF/image) pour ancrer le métier.",
    )
    add_bullet(
        doc,
        "Ne pas faire dépendre 100 % du message commercial de la 3D : le texte doit rester compréhensible seul.",
    )

    add_heading_custom(doc, "12.3 Contenu & crédibilité", 2)
    add_bullet(doc, "Fiche type projet : contexte, mission, livrables (plans), résultat, année.")
    add_bullet(doc, "Chronologie visuelle expérience (entreprise → freelance → collab. internationales).")
    add_bullet(doc, "Mention claire des domaines de plans couverts (à lister précisément avec le client).")
    add_bullet(doc, "Photo professionnelle / portrait discret : humanise la vitrine B2B.")

    add_heading_custom(doc, "12.4 Technique & exploitation", 2)
    add_bullet(doc, "Pipeline d’optimisation d’assets (gltf-transform) dans le repo.")
    add_bullet(doc, "Observabilité basique : taux d’arrivée section Contact, clics LinkedIn/mail.")
    add_bullet(doc, "Back-office minimal (phase 2) pour ajouter un projet sans redéployer du code.")
    add_bullet(doc, "Nom de domaine .fr si audience France ; e-mail pro sur le domaine.")

    add_heading_custom(doc, "12.5 Ce qu’il vaut mieux éviter", 2)
    add_bullet(doc, "Trop de sections « cartes » génériques type startup template.")
    add_bullet(doc, "Autoplay sonore, particules excessives, glow violet « IA default ».")
    add_bullet(doc, "Forcer la 3D haute qualité sur tous les téléphones bas de gamme.")
    add_bullet(doc, "Masquer les moyens de contact derrière trop d’animation.")

    # --- 13 ---
    add_heading_custom(doc, "13. Livrables et jalons", 1)
    add_table(
        doc,
        ["Jalon", "Contenu", "Résultat attendu"],
        [
            ("J0 — Cadrage", "CDC validé, storyboard, liste projets, contacts", "Scope figé MVP"),
            ("J1 — Design", "Direction artistique, maquettes desktop/mobile", "Validation visuelle"),
            ("J2 — Prototype 3D", "Maison + timeline scroll basique", "Preuve technique"),
            ("J3 — Intégration contenu", "Textes, projets, contact, légal", "Site navigable"),
            ("J4 — Polish & perf", "LOD, fallbacks, a11y, SEO", "Prêt préprod"),
            ("J5 — Mise en ligne", "Domaine, analytics, tests", "Production"),
        ],
    )

    add_heading_custom(doc, "13.1 Estimation indicative (ordre de grandeur)", 2)
    add_para(
        doc,
        "Selon la complexité du modèle 3D et la disponibilité des assets :\n"
        "• MVP (Next.js + 3D scroll + contenu) : 3 à 6 semaines.\n"
        "• Avec FastAPI + admin projets + i18n EN : +2 à 4 semaines.\n"
        "Ces durées dépendent fortement de la qualité/préparation du modèle 3D et des visuels projets.",
    )

    # --- 14 ---
    add_heading_custom(doc, "14. Critères d’acceptation", 1)
    add_bullet(doc, "Au chargement, la façade principale de la maison est visible et identifiable.")
    add_bullet(doc, "Le scroll fait progresser le zoom / l’approche de façon fluide et prévisible.")
    add_bullet(doc, "À un seuil défini, la porte s’ouvre ; une séquence de visite peut continuer sans scroll.")
    add_bullet(doc, "La visite conduit jusqu’au premier étage selon le storyboard validé.")
    add_bullet(doc, "Les contenus expérience / projets restent lisibles pendant le parcours.")
    add_bullet(doc, "En section Contact, la maison disparaît (ou s’efface) dynamiquement.")
    add_bullet(doc, "Liens LinkedIn, e-mail, Facebook fonctionnels.")
    add_bullet(doc, "Bouton skip / accès direct Contact disponible.")
    add_bullet(doc, "Fallback ou mode allégé si WebGL / perf insuffisante.")
    add_bullet(doc, "Pages légales présentes ; site déployé en HTTPS.")

    # --- 15 ---
    add_heading_custom(doc, "15. Annexes", 1)
    add_heading_custom(doc, "15.1 Décisions à valider avec le porteur de projet", 2)
    add_numbered(doc, "Liste exacte des types de plans proposés.")
    add_numbered(doc, "Sélection des 4–8 projets montrables (droits d’image / confidentialité).")
    add_numbered(doc, "URLs LinkedIn / Facebook / e-mail / téléphone définitifs.")
    add_numbered(doc, "Nom commercial affiché et slogan.")
    add_numbered(doc, "Faut-il un formulaire de contact dès le MVP ?")
    add_numbered(doc, "Budget / responsabilité de création du modèle 3D (interne, freelances 3D, achat).")
    add_numbered(doc, "Hébergement et nom de domaine.")

    add_heading_custom(doc, "15.2 Recommandation finale (synthèse technique)", 2)
    add_para(
        doc,
        "Pour ce repo et ce produit, la manière recommandée est :\n"
        "1) Construire le cœur en React/Next.js + React Three Fiber (expérience 3D scroll-driven).\n"
        "2) Garder le contenu et le SEO en HTML/React (pas uniquement dans le canvas).\n"
        "3) N’introduire Python (FastAPI) que si un vrai besoin backend apparaît (formulaire robuste, "
        "admin, traitements).\n"
        "4) Livrer un mode mobile dégradé dès le MVP pour ne pas sacrifier la conversion France mobile.\n"
        "5) Mesurer les clics Contact : c’est l’indicateur de succès, pas seulement la beauté de la 3D.",
    )

    add_para(
        doc,
        "Document généré pour le dépôt BTP-site- — Version 1.0 — Juillet 2026.",
        italic=True,
        size=10,
    )

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
